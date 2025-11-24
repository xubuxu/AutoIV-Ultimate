"""
AutoIV-Ultimate Reporting Module - Word Builder

Creates comprehensive Word reports with images, tables, and formatted text.
"""
import logging
from pathlib import Path
from typing import Dict, List, Optional

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


class WordBuilder:
    """Generates Word reports with images and tables."""
    
    def __init__(self, output_dir: str = "output"):
        """
        Initialize WordBuilder.
        
        Args:
            output_dir: Directory to save Word files
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        logger.info("WordBuilder initialized")
    
    def create_report(self, stats_df: pd.DataFrame, champion_df: pd.DataFrame,
                     img_paths: Dict[str, str], filename: str = "IV_Analysis_Report.docx") -> str:
        """
        Create comprehensive Word report.
        
        Args:
            stats_df: Statistics dataframe
            champion_df: Champion cells dataframe
            img_paths: Dictionary mapping plot names to file paths
            filename: Output filename
            
        Returns:
            Path to saved Word file
        """
        try:
            doc = Document()
            
            # Add title
            title = doc.add_heading('IV Analysis Report', level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add statistics section
            doc.add_heading('1. Statistical Summary', level=1)
            self._add_dataframe_table(doc, stats_df, "Batch Statistics")
            
            # Add champion cells section
            doc.add_heading('2. Champion Cells', level=1)
            self._add_dataframe_table(doc, champion_df, "Best Cell per Batch")
            
            # Add plots section
            doc.add_heading('3. Visualizations', level=1)
            self._add_plots(doc, img_paths)
            
            # Save
            output_path = self.output_dir / filename
            doc.save(output_path)
            
            logger.info(f"Word report saved: {output_path}")
            return str(output_path)
            
        except Exception as e:
            logger.error(f"Word report generation failed: {e}")
            return ""
    
    def _add_dataframe_table(self, doc: Document, df: pd.DataFrame, caption: str):
        """Add a pandas DataFrame as a formatted table."""
        if df.empty:
            doc.add_paragraph(f"{caption}: No data available")
            return
        
        # Add caption
        p = doc.add_paragraph(caption)
        p.runs[0].bold = True
        
        # Create table
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Light Grid Accent 1'
        
        # Add headers
        hdr_cells = table.rows[0].cells
        for i, column in enumerate(df.columns):
            hdr_cells[i].text = str(column)
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Add data rows
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)
        
        doc.add_paragraph()  # Add spacing
    
    def _add_plots(self, doc: Document, img_paths: Dict[str, str]):
        """Add plot images to document."""
        plot_titles = {
            'boxplot': 'Parameter Boxplots',
            'histogram': 'Efficiency Distribution',
            'trend': 'Batch Trends',
            'yield': 'Yield Distribution',
            'jv_curves': 'Champion J-V Curves',
            'voc_jsc': 'Voc vs Jsc Correlation',
            'drivers': 'Efficiency Drivers',
            'resistance': 'Resistance Analysis',
            'corr_matrix': 'Correlation Matrix'
        }
        
        for key, path in img_paths.items():
            if path and Path(path).exists():
                title = plot_titles.get(key, key.replace('_', ' ').title())
                
                doc.add_heading(title, level=2)
                try:
                    doc.add_picture(path, width=Inches(6))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    logger.warning(f"Failed to add image {path}: {e}")
                    doc.add_paragraph(f"[Image not available: {path}]")
                
                doc.add_paragraph()  # Add spacing
